import sys
import traceback
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QPushButton, QLabel, QVBoxLayout,
    QHBoxLayout, QWidget, QListWidget, QListWidgetItem, QTableWidget, QTableWidgetItem,
    QHeaderView, QMessageBox, QLineEdit, QProgressBar
)
from PyQt5.QtGui import QIcon, QFontMetrics
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

class FileListItem(QWidget):
    def __init__(self, file_path, remove_callback):
        super().__init__()
        self.file_path = file_path
        self.remove_callback = remove_callback

        layout = QHBoxLayout()
        self.setLayout(layout)

        self.label = QLabel(file_path)
        layout.addWidget(self.label)

        self.remove_button = QPushButton("❌")
        self.remove_button.setFixedSize(20, 20)
        self.remove_button.clicked.connect(self.remove_item)
        layout.addWidget(self.remove_button)

    def remove_item(self):
        self.remove_callback(self.file_path)

class SaveChangesThread(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal()
    error = pyqtSignal(str)
    update_status = pyqtSignal(str)  # Новый сигнал для обновления статуса

    def __init__(self, table, files, parent=None):
        super(SaveChangesThread, self).__init__(parent)
        self.table = table
        self.files = files

    def run(self):
        try:
            row_count = self.table.rowCount()
            row_idx = 0
            for row in range(row_count):
                original_text = self.table.item(row, 0).text()
                file_locations = self.table.item(row, 1).text().split("\n")
                new_text = self.table.item(row, 2).text()
                replacement_count = 0  # Для подсчета замен

                if new_text and original_text != new_text:
                    for file_index, file_location in enumerate(file_locations):
                        file_name, location = file_location.split(" (")
                        location = location.rstrip(")")
                        file_path = next((file for file in self.files if file.endswith(file_name)), None)

                        if file_path:
                            try:
                                file_replacements = 0  # Подсчет замен в текущем файле
                                
                                if file_path.endswith('.xlsx'):
                                    try:
                                        workbook = load_workbook(file_path)
                                    except Exception as e:
                                        self.error.emit(f"Ошибка при открытии Excel-файла {file_path}: {str(e)}")
                                        continue
                                    
                                    worksheet = workbook[location]
                                    for row_cells in worksheet.iter_rows():
                                        for cell in row_cells:
                                            if cell.value and str(cell.value) == original_text:
                                                cell.value = new_text
                                                cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
                                                file_replacements += 1
                                    workbook.save(file_path)
                                elif file_path.endswith('.docx'):
                                    try:
                                        doc = Document(file_path)
                                    except Exception as e:
                                        self.error.emit(f"Ошибка при открытии Word-файла {file_path}: {str(e)}")
                                        continue
                                    
                                    for table in doc.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                if cell.text.strip() == original_text:
                                                    cell.text = new_text
                                                    for paragraph in cell.paragraphs:
                                                        for run in paragraph.runs:
                                                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                                    file_replacements += 1
                                    doc.save(file_path)

                                # Обновляем статус
                                replacement_count += file_replacements
                                self.update_status.emit(f"Файл: {file_path}, Замена: {replacement_count}, Номер вхождения: {file_index + 1}")
                            except Exception as e:
                                self.error.emit(f"Ошибка при обработке файла {file_path}: {str(e)}")
                                return

                row_idx += 1
                self.progress.emit(int(row_idx / row_count * 100))
            self.finished.emit()
        except Exception as e:
            self.error.emit(f"Общая ошибка: {str(e)}")
            traceback.print_exc()

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Замена текста в файлах Word и Excel")
        self.setGeometry(100, 100, 1000, 600)

        self.layout = QVBoxLayout()

        self.file_button = QPushButton("Добавить файлы")
        self.file_button.clicked.connect(self.open_file_dialog)
        self.layout.addWidget(self.file_button)

        self.file_list = QListWidget()
        self.layout.addWidget(self.file_list)

        self.scan_button = QPushButton("Сканировать файлы")
        self.scan_button.clicked.connect(self.scan_files)
        self.layout.addWidget(self.scan_button)

        self.search_bar = QLineEdit()
        self.search_bar.setPlaceholderText("Поиск полей...")
        self.search_bar.textChanged.connect(self.filter_matches)
        self.layout.addWidget(self.search_bar)

        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["Повторяющееся значение", "Файлы и местоположения", "Новое слово", "Количество"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setSortingEnabled(True)  # Сортировка по столбцам
        self.layout.addWidget(self.table)

        self.save_button = QPushButton("Сохранить изменения")
        self.save_button.clicked.connect(self.start_saving_changes)
        self.layout.addWidget(self.save_button)
        self.save_button.setEnabled(False)

        self.progress_bar = QProgressBar()
        self.layout.addWidget(self.progress_bar)

        self.status_label = QLabel("Статус: Готов")
        self.layout.addWidget(self.status_label)

        container = QWidget()
        container.setLayout(self.layout)
        self.setCentralWidget(container)

        self.files = set()
        self.matches = {}

    def open_file_dialog(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Выберите файлы", "", "Excel and Word Files (*.xlsx *.docx)")
        if files:
            self.files.update(files)  # Используем set для исключения дубликатов
            self.update_file_list()

    def update_file_list(self):
        self.file_list.clear()
        for file in self.files:
            item_widget = FileListItem(file, self.remove_file)
            list_item = QListWidgetItem(self.file_list)
            list_item.setSizeHint(item_widget.sizeHint())
            self.file_list.addItem(list_item)
            self.file_list.setItemWidget(list_item, item_widget)

    def remove_file(self, file_path):
        if file_path in self.files:
            self.files.remove(file_path)
            self.update_file_list()

    def scan_files(self):
        self.status_label.setText("Статус: Сканирование файлов...")
        self.matches = {}
        unique_values = {}

        for file in self.files:
            file_name = file.split('/')[-1]
            if file.endswith('.xlsx'):
                try:
                    workbook = load_workbook(file)
                    for sheet in workbook.sheetnames:
                        worksheet = workbook[sheet]
                        for row in worksheet.iter_rows(values_only=True):
                            for cell in row:
                                if cell:
                                    value = str(cell)
                                    if value in unique_values:
                                        unique_values[value].append((file_name, sheet))
                                    else:
                                        unique_values[value] = [(file_name, sheet)]
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии Excel-файла {file}: {str(e)}")
                    continue
            elif file.endswith('.docx'):
                try:
                    doc = Document(file)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                value = cell.text.strip()
                                if value:
                                    if value in unique_values:
                                        unique_values[value].append((file_name, "Таблица"))
                                    else:
                                        unique_values[value] = [(file_name, "Таблица")]
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", f"Ошибка при открытии Word-файла {file}: {str(e)}")
                    continue

        self.matches = {k: v for k, v in unique_values.items() if len(v) > 1}
        self.display_matches()
        self.status_label.setText("Статус: Сканирование завершено")


    def filter_matches(self):
        filter_text = self.search_bar.text().lower()
        for row in range(self.table.rowCount()):
            item = self.table.item(row, 0)
            self.table.setRowHidden(row, filter_text not in item.text().lower())

    def display_matches(self):
        self.table.setRowCount(0)
        if not self.matches:
            QMessageBox.information(self, "Результат сканирования", "Совпадений не найдено.")
        else:
            for match, locations in self.matches.items():
                row_position = self.table.rowCount()
                self.table.insertRow(row_position)

                item_text = QTableWidgetItem(match)
                item_text.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.table.setItem(row_position, 0, item_text)

                unique_locations = set()

                for file, location in locations:
                    unique_locations.add((file, location))

                file_locations = "\n".join([f"{file} ({location})" for file, location in unique_locations])

                item_file_locations = QTableWidgetItem(file_locations)
                item_file_locations.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled)
                self.table.setItem(row_position, 1, item_file_locations)

                new_word_item = QTableWidgetItem(match)
                self.table.setItem(row_position, 2, new_word_item)

                # Добавляем количество вхождений
                count_item = QTableWidgetItem(str(len(unique_locations)))
                count_item.setFlags(Qt.ItemIsSelectable | Qt.ItemIsEnabled | Qt.ItemIsDragEnabled | Qt.ItemIsDropEnabled)
                self.table.setItem(row_position, 3, count_item)

                font_metrics = QFontMetrics(self.table.font())
                text_height = font_metrics.lineSpacing() * (file_locations.count('\n') + 1)
                self.table.setRowHeight(row_position, text_height + 10)

            self.save_button.setEnabled(True)

    def start_saving_changes(self):
        self.save_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.status_label.setText("Статус: Сохранение изменений...")
        self.save_thread = SaveChangesThread(self.table, self.files)
        self.save_thread.progress.connect(self.progress_bar.setValue)
        self.save_thread.finished.connect(self.on_save_finished)
        self.save_thread.error.connect(self.on_save_error)
        self.save_thread.update_status.connect(self.status_label.setText)  # Подключаем сигнал обновления статуса
        self.save_thread.start()


    def on_save_finished(self):
        QMessageBox.information(self, "Успех", "Изменения успешно сохранены.")
        self.save_button.setEnabled(True)
        self.status_label.setText("Статус: Завершено")

    def on_save_error(self, error_message):
        QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при сохранении файлов: {error_message}")
        self.save_button.setEnabled(True)
        self.status_label.setText("Статус: Ошибка")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
